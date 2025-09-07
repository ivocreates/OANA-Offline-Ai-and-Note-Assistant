"""
DOCX Parser utility for extracting text from Word documents
"""

import os
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("python-docx not available. DOCX parsing will be limited.")

class DocxParser:
    def __init__(self):
        self.available = PYTHON_DOCX_AVAILABLE
        
    def is_available(self):
        """Check if DOCX parsing is available"""
        return self.available
        
    def extract_text(self, filepath):
        """Extract text from DOCX file"""
        if not self.available:
            raise Exception("python-docx not installed. Install with: pip install python-docx")
            
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"DOCX file not found: {filepath}")
            
        try:
            # Handle both .docx and .doc files
            if filepath.lower().endswith('.doc') and not filepath.lower().endswith('.docx'):
                return self._extract_from_doc(filepath)
            else:
                return self._extract_from_docx(filepath)
                
        except Exception as e:
            raise Exception(f"Failed to extract text from document: {str(e)}")
            
    def _extract_from_docx(self, filepath):
        """Extract text from DOCX file"""
        doc = Document(filepath)
        text_content = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content += cell.text + " "
                text_content += "\n"
                
        # Clean up text
        text_content = self._clean_text(text_content)
        
        return text_content
        
    def _extract_from_doc(self, filepath):
        """Extract text from .doc file (legacy format)"""
        # For .doc files, we can try using python-docx2txt or fallback methods
        try:
            import docx2txt
            text_content = docx2txt.process(filepath)
            return self._clean_text(text_content)
        except ImportError:
            # Fallback: try to read as text (may not work well)
            try:
                with open(filepath, 'rb') as file:
                    raw_text = file.read().decode('utf-8', errors='ignore')
                    # Basic cleanup for binary content
                    text_lines = []
                    for line in raw_text.split('\n'):
                        # Filter out lines with too many non-printable characters
                        printable_chars = sum(1 for c in line if c.isprintable())
                        if len(line) > 0 and printable_chars / len(line) > 0.7:
                            text_lines.append(line.strip())
                    
                    return '\n'.join(text_lines)
            except:
                raise Exception("Cannot process .doc files. Please convert to .docx format or install docx2txt")
                
    def extract_metadata(self, filepath):
        """Extract metadata from DOCX file"""
        if not self.available:
            return {}
            
        try:
            doc = Document(filepath)
            core_props = doc.core_properties
            
            return {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision or 0
            }
            
        except Exception as e:
            print(f"Failed to extract DOCX metadata: {e}")
            return {}
            
    def _clean_text(self, text):
        """Clean extracted text"""
        if not text:
            return ""
            
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
                
        # Join lines with proper spacing
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive newlines
        while '\n\n\n' in cleaned_text:
            cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
            
        return cleaned_text
        
    def get_paragraph_count(self, filepath):
        """Get number of paragraphs in document"""
        if not self.available:
            return 0
            
        try:
            doc = Document(filepath)
            return len([p for p in doc.paragraphs if p.text.strip()])
        except:
            return 0
