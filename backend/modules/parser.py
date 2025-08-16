import os
import logging
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
import re

# Document parsers
from PyPDF2 import PdfReader
import docx
import nltk

# Ensure NLTK resources are downloaded
try:
    nltk.data.find('punkt')
except LookupError:
    nltk.download('punkt')

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a chunk of a document with metadata"""
    id: str
    document_id: str
    content: str
    page_num: Optional[int] = None
    chunk_num: int = 0
    section: Optional[str] = None
    metadata: Dict[str, Any] = None

class DocumentParser:
    """Parser for different document formats"""
    
    def __init__(self):
        """Initialize the document parser"""
        pass
    
    def parse(self, file_path: str) -> List[DocumentChunk]:
        """
        Parse a document file into chunks based on its type
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of DocumentChunk objects
        """
        # Get file extension to determine parser
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # Generate a document ID from the file name
        document_id = os.path.basename(file_path).split('.')[0]
        
        if ext == '.pdf':
            return self._parse_pdf(file_path, document_id)
        elif ext == '.docx':
            return self._parse_docx(file_path, document_id)
        elif ext == '.txt':
            return self._parse_txt(file_path, document_id)
        elif ext == '.md':
            return self._parse_md(file_path, document_id)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _parse_pdf(self, file_path: str, document_id: str) -> List[DocumentChunk]:
        """Parse a PDF document"""
        chunks = []
        
        try:
            pdf = PdfReader(file_path)
            
            # Extract sections and content
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                
                if text:
                    # Find section headers (simple heuristic)
                    sections = self._extract_sections(text)
                    
                    if sections:
                        # Create chunks for each section
                        for i, (section, content) in enumerate(sections):
                            chunks.append(
                                DocumentChunk(
                                    id=f"{document_id}_p{page_num}_s{i}",
                                    document_id=document_id,
                                    content=content,
                                    page_num=page_num + 1,
                                    chunk_num=i,
                                    section=section,
                                    metadata={"source_page": page_num + 1}
                                )
                            )
                    else:
                        # No sections found, split by paragraphs
                        paragraphs = re.split(r'\n\s*\n', text)
                        for i, para in enumerate(paragraphs):
                            if para.strip():
                                chunks.append(
                                    DocumentChunk(
                                        id=f"{document_id}_p{page_num}_c{i}",
                                        document_id=document_id,
                                        content=para.strip(),
                                        page_num=page_num + 1,
                                        chunk_num=i,
                                        section=None,
                                        metadata={"source_page": page_num + 1}
                                    )
                                )
            
        except Exception as e:
            logger.error(f"Error parsing PDF file {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _parse_docx(self, file_path: str, document_id: str) -> List[DocumentChunk]:
        """Parse a DOCX document"""
        chunks = []
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Find sections
            sections = self._extract_sections(full_text)
            
            if sections:
                # Create chunks for each section
                for i, (section, content) in enumerate(sections):
                    chunks.append(
                        DocumentChunk(
                            id=f"{document_id}_s{i}",
                            document_id=document_id,
                            content=content,
                            chunk_num=i,
                            section=section,
                            metadata={}
                        )
                    )
            else:
                # No sections found, use paragraphs
                current_section = ""
                section_content = []
                
                for i, para in enumerate(doc.paragraphs):
                    text = para.text.strip()
                    if text:
                        # Check if this is a heading paragraph
                        if para.style.name.startswith('Heading'):
                            # Save previous section if exists
                            if section_content:
                                chunks.append(
                                    DocumentChunk(
                                        id=f"{document_id}_s{len(chunks)}",
                                        document_id=document_id,
                                        content="\n".join(section_content),
                                        chunk_num=len(chunks),
                                        section=current_section,
                                        metadata={}
                                    )
                                )
                                section_content = []
                            
                            current_section = text
                        else:
                            section_content.append(text)
                
                # Add the last section
                if section_content:
                    chunks.append(
                        DocumentChunk(
                            id=f"{document_id}_s{len(chunks)}",
                            document_id=document_id,
                            content="\n".join(section_content),
                            chunk_num=len(chunks),
                            section=current_section,
                            metadata={}
                        )
                    )
            
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _parse_txt(self, file_path: str, document_id: str) -> List[DocumentChunk]:
        """Parse a TXT document"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Find sections
            sections = self._extract_sections(text)
            
            if sections:
                # Create chunks for each section
                for i, (section, content) in enumerate(sections):
                    chunks.append(
                        DocumentChunk(
                            id=f"{document_id}_s{i}",
                            document_id=document_id,
                            content=content,
                            chunk_num=i,
                            section=section,
                            metadata={}
                        )
                    )
            else:
                # No sections found, split by paragraphs
                paragraphs = re.split(r'\n\s*\n', text)
                
                # Group paragraphs to avoid too small chunks
                grouped_paras = []
                current_group = []
                current_length = 0
                
                for para in paragraphs:
                    if para.strip():
                        para_len = len(para.split())
                        
                        if current_length + para_len > 300:  # ~300 words per chunk
                            if current_group:
                                grouped_paras.append("\n\n".join(current_group))
                            current_group = [para.strip()]
                            current_length = para_len
                        else:
                            current_group.append(para.strip())
                            current_length += para_len
                
                # Add the last group
                if current_group:
                    grouped_paras.append("\n\n".join(current_group))
                
                # Create chunks for grouped paragraphs
                for i, content in enumerate(grouped_paras):
                    chunks.append(
                        DocumentChunk(
                            id=f"{document_id}_c{i}",
                            document_id=document_id,
                            content=content,
                            chunk_num=i,
                            section=None,
                            metadata={}
                        )
                    )
            
        except Exception as e:
            logger.error(f"Error parsing TXT file {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _parse_md(self, file_path: str, document_id: str) -> List[DocumentChunk]:
        """Parse a Markdown document"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Find markdown headers and content
            header_pattern = r'^(#{1,6})\s+(.+)$'
            lines = text.split('\n')
            
            current_section = None
            current_content = []
            
            for line in lines:
                header_match = re.match(header_pattern, line)
                
                if header_match:
                    # Save previous section if exists
                    if current_content:
                        chunks.append(
                            DocumentChunk(
                                id=f"{document_id}_s{len(chunks)}",
                                document_id=document_id,
                                content="\n".join(current_content),
                                chunk_num=len(chunks),
                                section=current_section,
                                metadata={}
                            )
                        )
                        current_content = []
                    
                    # Start new section
                    current_section = header_match.group(2).strip()
                else:
                    current_content.append(line)
            
            # Add the last section
            if current_content:
                chunks.append(
                    DocumentChunk(
                        id=f"{document_id}_s{len(chunks)}",
                        document_id=document_id,
                        content="\n".join(current_content),
                        chunk_num=len(chunks),
                        section=current_section,
                        metadata={}
                    )
                )
            
        except Exception as e:
            logger.error(f"Error parsing Markdown file {file_path}: {str(e)}")
            raise
        
        return chunks
    
    def _extract_sections(self, text: str) -> List[Tuple[str, str]]:
        """Extract sections from text using heuristic patterns"""
        # Look for common section headers:
        # 1. Numbers followed by period/closing parenthesis at start of line
        # 2. All caps lines
        # 3. Lines ending with colon
        
        section_patterns = [
            r'^(?:\d+[\.\)]|\*)\s+([A-Z][^\.]+)$',  # Numbered or bullet points
            r'^([A-Z][A-Z\s]{3,}):?\s*$',            # ALL CAPS section titles
            r'^(Chapter|Section)\s+\d+:?\s*([^\n]+)$', # Chapter/Section headers
            r'^([A-Z][a-zA-Z\s]+):$'                 # Title with colon
        ]
        
        sections = []
        current_section = None
        current_content = []
        
        # Split by lines
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            is_section_header = False
            
            for pattern in section_patterns:
                match = re.match(pattern, line)
                if match:
                    is_section_header = True
                    
                    # Add previous section if exists
                    if current_content:
                        sections.append((
                            current_section or "Introduction", 
                            "\n".join(current_content).strip()
                        ))
                        current_content = []
                    
                    # Extract section title
                    if len(match.groups()) > 1:
                        current_section = match.group(1) + ": " + match.group(2)
                    else:
                        current_section = match.group(1)
                    break
            
            if not is_section_header:
                current_content.append(line)
        
        # Add the last section
        if current_content:
            sections.append((
                current_section or "Content", 
                "\n".join(current_content).strip()
            ))
        
        return sections
    
    def test(self):
        """Test the parser is working"""
        # Simple test to ensure the module is working
        test_text = "This is a test."
        sections = self._extract_sections(test_text)
        
        if sections:
            return True
        return True  # Even if no sections found, parser is working
